#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
* @Author h3st4k3r
* @Version Mark.1
* With this class you will be able to generate the latest monthly report based on your MISP. 
* You will have to change the path where you want to save it.
* GNU General Public License
* This script is part of Hegemon's shadow project
*
"""

from pymisp import ExpandedPyMISP
from docx import Document
import datetime

class Reporting:
    def __init__(self, misp_url, misp_key, count, misp_verifycert=False):
        self.misp_url = misp_url
        self.misp_key = misp_key
        self.misp_verifycert = misp_verifycert
        self.count = count

    def get_lastEvents(self):
        misp = ExpandedPyMISP(self.misp_url, self.misp_key, self.misp_verifycert)
        response = misp.direct_call('events/', None)
        return response


    def generate_wordDocument(self, events):
        current_date = datetime.date.today()
        one_month_ago = current_date - datetime.timedelta(days=30)

        document = Document()

        document.add_heading('REPORTE MENSUAL', level=1)
        document.add_heading('Eventos desde '+str(one_month_ago)+' hasta ' + str(current_date), level=2)
        document.add_heading("By Hegemon's Shadow", level=3)
        document.add_paragraph('')
        for event in events:
            event_date = datetime.datetime.strptime(event['date'], "%Y-%m-%d").date()

            if one_month_ago <= event_date <= current_date:
                document.add_paragraph(f'ID: {str(event["id"])}')
                document.add_paragraph(f'Fecha: {event["date"]}')
                document.add_paragraph(f'Información: {event["info"]}')
                document.add_paragraph(f'ID único: {event["uuid"]}')
                tags = [tag['Tag']['name'] for tag in event['EventTag']]
                tags_str = ', '.join(tags)
                document.add_paragraph(f'Etiquetas: {tags_str}')
                document.add_paragraph('---')

        document.save('./reports/HegemonsShadowReport'+str(current_date)+'.docx')

    def run(self):
        top_events = self.get_lastEvents()

        if top_events:
            self.generate_wordDocument(top_events)
        else:
            print("No se pudo obtener los eventos de MISP.")
